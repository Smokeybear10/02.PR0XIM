from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from io import BytesIO
import traceback


class ResumeBuilder:
    def __init__(self):
        self.templates = {
            "modern": self.build_modern_template,
            "professional": self.build_professional_template,
            "minimal": self.build_minimal_template,
            "creative": self.build_creative_template,
        }

    def generate_resume(self, data: dict) -> BytesIO:
        try:
            doc = Document()
            template_name = data.get("template", "modern").lower()
            builder = self.templates.get(template_name, self.build_modern_template)
            doc = builder(doc, data)

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer
        except Exception as e:
            print(f"Error in generate_resume: {str(e)}")
            print(traceback.format_exc())
            raise

    def _format_list_items(self, items) -> list[str]:
        if isinstance(items, str):
            return [item.strip() for item in items.split("\n") if item.strip()]
        if isinstance(items, list):
            return [item.strip() for item in items if item and item.strip()]
        return []

    def _add_contact_block(self, doc, data: dict, style):
        contact_parts = []
        pi = data.get("personal_info", {})
        if pi.get("email"):
            contact_parts.append(pi["email"])
        if pi.get("phone"):
            contact_parts.append(pi["phone"])
        if pi.get("location"):
            contact_parts.append(pi["location"])
        if contact_parts:
            p = doc.add_paragraph()
            p.style = style
            p.add_run(" | ".join(contact_parts))

        links = []
        if pi.get("linkedin"):
            links.append(f"LinkedIn: {pi['linkedin']}")
        if pi.get("portfolio"):
            links.append(f"Portfolio: {pi['portfolio']}")
        if links:
            p = doc.add_paragraph()
            p.style = style
            p.add_run(" | ".join(links))

    def _add_experience_section(self, doc, data: dict, section_style, normal_style):
        if not data.get("experience"):
            return
        doc.add_paragraph("EXPERIENCE", style=section_style)
        for exp in data["experience"]:
            p = doc.add_paragraph()
            p.style = normal_style
            p.add_run(f"{exp['position']} at {exp['company']}").bold = True
            p.add_run(f"\n{exp['start_date']} - {exp['end_date']}")
            if exp.get("description"):
                desc = doc.add_paragraph(exp["description"])
                desc.style = normal_style
                desc.paragraph_format.left_indent = Inches(0.3)
            if exp.get("responsibilities"):
                for resp in self._format_list_items(exp["responsibilities"]):
                    bullet = doc.add_paragraph()
                    bullet.style = normal_style
                    bullet.paragraph_format.left_indent = Inches(0.5)
                    bullet.add_run("* " + resp)

    def _add_projects_section(self, doc, data: dict, section_style, normal_style):
        if not data.get("projects"):
            return
        doc.add_paragraph("PROJECTS", style=section_style)
        for proj in data["projects"]:
            p = doc.add_paragraph()
            p.style = normal_style
            p.add_run(proj["name"]).bold = True
            if proj.get("technologies"):
                p.add_run(f" | {proj['technologies']}")
            if proj.get("description"):
                desc = doc.add_paragraph(proj["description"])
                desc.style = normal_style
                desc.paragraph_format.left_indent = Inches(0.3)
            if proj.get("responsibilities"):
                for resp in self._format_list_items(proj["responsibilities"]):
                    bullet = doc.add_paragraph()
                    bullet.style = normal_style
                    bullet.paragraph_format.left_indent = Inches(0.5)
                    bullet.add_run("* " + resp)

    def _add_education_section(self, doc, data: dict, section_style, normal_style):
        if not data.get("education"):
            return
        doc.add_paragraph("EDUCATION", style=section_style)
        for edu in data["education"]:
            p = doc.add_paragraph()
            p.style = normal_style
            p.add_run(f"{edu['school']}").bold = True
            p.add_run(f"\n{edu['degree']} in {edu['field']}")
            p.add_run(f"\nGraduation: {edu['graduation_date']}")
            if edu.get("gpa"):
                p.add_run(f" | GPA: {edu['gpa']}")

    def _add_skills_section(self, doc, data: dict, section_style, normal_style):
        if not data.get("skills"):
            return
        doc.add_paragraph("SKILLS", style=section_style)
        skills = data["skills"]
        for category_key, title in [
            ("technical", "Technical Skills"),
            ("soft", "Soft Skills"),
            ("languages", "Languages"),
            ("tools", "Tools & Technologies"),
        ]:
            items = skills.get(category_key)
            if items:
                p = doc.add_paragraph()
                p.style = normal_style
                p.add_run(f"{title}: ").bold = True
                p.add_run(" | ".join(self._format_list_items(items)))

    def _get_or_create_style(self, styles, name: str, font_size: int, bold: bool = False,
                             color: RGBColor | None = None, font_name: str = "Arial",
                             alignment=None, space_before: int = 0, space_after: int = 2):
        if name not in styles:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = styles[name]
        style.font.size = Pt(font_size)
        style.font.bold = bold
        style.font.name = font_name
        if color:
            style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(space_before)
        style.paragraph_format.space_after = Pt(space_after)
        if alignment is not None:
            style.paragraph_format.alignment = alignment
        return style

    def build_modern_template(self, doc, data: dict):
        styles = doc.styles
        name_style = self._get_or_create_style(
            styles, "Modern Name", 24, bold=True,
            color=RGBColor(41, 128, 185), alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=6, space_after=0,
        )
        section_style = self._get_or_create_style(
            styles, "Modern Section", 14, bold=True,
            color=RGBColor(41, 128, 185), space_before=16, space_after=4,
        )
        normal_style = self._get_or_create_style(
            styles, "Modern Normal", 10, color=RGBColor(44, 62, 80),
        )
        contact_style = self._get_or_create_style(
            styles, "Modern Contact", 10,
            color=RGBColor(41, 128, 185), alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )

        p = doc.add_paragraph(data["personal_info"]["full_name"].upper())
        p.style = name_style

        self._add_contact_block(doc, data, contact_style)

        if data.get("summary"):
            doc.add_paragraph("PROFESSIONAL SUMMARY", style=section_style)
            s = doc.add_paragraph(data["summary"])
            s.style = normal_style

        self._add_experience_section(doc, data, section_style, normal_style)
        self._add_projects_section(doc, data, section_style, normal_style)
        self._add_education_section(doc, data, section_style, normal_style)
        self._add_skills_section(doc, data, section_style, normal_style)

        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

        return doc

    def build_professional_template(self, doc, data: dict):
        styles = doc.styles
        header_style = self._get_or_create_style(
            styles, "Pro Header", 24, bold=True,
            color=RGBColor(0, 0, 0), font_name="Calibri", space_after=4,
        )
        section_style = self._get_or_create_style(
            styles, "Pro Section", 14, bold=True,
            color=RGBColor(0, 120, 215), font_name="Calibri", space_before=12, space_after=6,
        )
        normal_style = self._get_or_create_style(
            styles, "Pro Normal", 10, font_name="Calibri",
        )
        contact_style = self._get_or_create_style(
            styles, "Pro Contact", 10, font_name="Calibri", space_after=6,
        )

        p = doc.add_paragraph(data["personal_info"]["full_name"])
        p.style = header_style

        self._add_contact_block(doc, data, contact_style)

        if data.get("summary"):
            doc.add_paragraph("PROFESSIONAL SUMMARY", style=section_style)
            s = doc.add_paragraph(data["summary"])
            s.style = normal_style

        self._add_experience_section(doc, data, section_style, normal_style)
        self._add_projects_section(doc, data, section_style, normal_style)
        self._add_education_section(doc, data, section_style, normal_style)
        self._add_skills_section(doc, data, section_style, normal_style)
        return doc

    def build_minimal_template(self, doc, data: dict):
        styles = doc.styles
        name_style = self._get_or_create_style(
            styles, "Min Name", 20, bold=True,
            color=RGBColor(50, 50, 50), space_after=2,
        )
        section_style = self._get_or_create_style(
            styles, "Min Section", 12, bold=True,
            color=RGBColor(80, 80, 80), space_before=10, space_after=4,
        )
        normal_style = self._get_or_create_style(
            styles, "Min Normal", 10, color=RGBColor(60, 60, 60),
        )
        contact_style = self._get_or_create_style(
            styles, "Min Contact", 9, color=RGBColor(100, 100, 100),
        )

        p = doc.add_paragraph(data["personal_info"]["full_name"])
        p.style = name_style

        self._add_contact_block(doc, data, contact_style)

        if data.get("summary"):
            doc.add_paragraph("Summary", style=section_style)
            s = doc.add_paragraph(data["summary"])
            s.style = normal_style

        self._add_experience_section(doc, data, section_style, normal_style)
        self._add_education_section(doc, data, section_style, normal_style)
        self._add_skills_section(doc, data, section_style, normal_style)
        return doc

    def build_creative_template(self, doc, data: dict):
        styles = doc.styles
        name_style = self._get_or_create_style(
            styles, "Cre Name", 26, bold=True,
            color=RGBColor(231, 76, 60), alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=0,
        )
        section_style = self._get_or_create_style(
            styles, "Cre Section", 14, bold=True,
            color=RGBColor(231, 76, 60), space_before=14, space_after=4,
        )
        normal_style = self._get_or_create_style(
            styles, "Cre Normal", 10, color=RGBColor(44, 62, 80),
        )
        contact_style = self._get_or_create_style(
            styles, "Cre Contact", 10,
            color=RGBColor(231, 76, 60), alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )

        p = doc.add_paragraph(data["personal_info"]["full_name"].upper())
        p.style = name_style

        self._add_contact_block(doc, data, contact_style)

        if data.get("summary"):
            doc.add_paragraph("ABOUT ME", style=section_style)
            s = doc.add_paragraph(data["summary"])
            s.style = normal_style

        self._add_experience_section(doc, data, section_style, normal_style)
        self._add_projects_section(doc, data, section_style, normal_style)
        self._add_education_section(doc, data, section_style, normal_style)
        self._add_skills_section(doc, data, section_style, normal_style)
        return doc
